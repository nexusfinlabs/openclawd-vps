#!/usr/bin/env python3
"""
Document Generator CLI — Template-based PDF + DOCX generation.

Usage:
    python document_generator_cli.py <DOC_TYPE> <CONTEXT>

Examples:
    python document_generator_cli.py NDA "empresa: Stripe, contacto: John Smith, jurisdicción: España"
    python document_generator_cli.py SOW "cliente: BBVA, scope: AI advisory, duración: 3 meses, fee: 15K"
    python document_generator_cli.py PROPUESTA "empresa: Mercado Pago, contexto: M&A advisory LatAm"

If a template exists for the doc_type (nda.html, sow.html, propuesta.html),
the LLM expands the brief context into structured variables for the template.
Otherwise, falls back to generic document generation.
"""

import os
import sys
import json
import uuid
import requests
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── Config ───────────────────────────────────────────
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
OUTPUT_DIR = "/docs/drafts"
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"

# Template-aware doc types
TEMPLATE_MAP = {
    "NDA": "nda.html",
    "SOW": "sow.html",
    "PROPUESTA": "propuesta.html",
    "PROPOSAL": "propuesta.html",
}

# LLM prompts per doc type for context expansion
EXPANSION_PROMPTS = {
    "NDA": """Eres un abogado corporativo experto. A partir del siguiente contexto breve, genera un JSON con los campos necesarios para rellenar un NDA profesional.

Contexto del usuario: "{context}"

Genera un JSON con estos campos (usa valores por defecto inteligentes si faltan):
- parte_reveladora: nombre de la empresa (default: "Nexus FinLabs S.L.")
- direccion_reveladora: dirección (default: "Madrid, España")
- parte_receptora: nombre de la otra parte
- direccion_receptora: dirección de la otra parte
- proposito: propósito del NDA
- vigencia: período de vigencia (default: "2 años")
- jurisdiccion: país de jurisdicción (default: "España")
- ciudad_jurisdiccion: ciudad para tribunales (default: "Madrid")
- info_adicional: información confidencial adicional específica al contexto

Responde ÚNICAMENTE con el JSON, sin markdown ni explicaciones.""",

    "SOW": """Eres un consultor de management experto en statements of work. A partir del siguiente contexto breve, genera un JSON con los campos necesarios para un SOW profesional.

Contexto del usuario: "{context}"

Genera un JSON con estos campos (usa valores por defecto inteligentes si faltan):
- cliente: nombre del cliente
- proveedor: nombre del proveedor (default: "Nexus FinLabs / iAgrowth")
- antecedentes: párrafo de antecedentes
- alcance_resumen: descripción del alcance del trabajo
- fases: lista de objetos con {{nombre, descripcion, duracion, entregables: [lista de strings]}}
- duracion: duración total del proyecto
- fee: fee total
- condiciones_pago: condiciones de pago (default: "50% al inicio, 50% a la entrega final")
- responsable_proveedor: responsable del proveedor (default: "Alberto Lebrón — Tech Advisor")
- contacto_cliente: contacto del cliente (default: "A designar")

Responde ÚNICAMENTE con el JSON, sin markdown ni explicaciones.""",

    "PROPUESTA": """Eres un consultor de negocio senior experto en propuestas comerciales. A partir del siguiente contexto breve, genera un JSON con los campos necesarios para una propuesta comercial profesional y convincente.

Contexto del usuario: "{context}"

Genera un JSON con estos campos (usa valores por defecto inteligentes si faltan):
- empresa_destino: nombre de la empresa destino
- preparado_por: quién prepara la propuesta (default: "Alberto Lebrón — Nexus FinLabs")
- resumen_ejecutivo: párrafo de resumen ejecutivo (compelling, orientado a resultados)
- contexto: párrafo de contexto y oportunidad
- solucion: descripción de la solución propuesta
- servicios: lista de objetos con {{nombre, descripcion}}
- valor_items: lista de strings con el valor añadido
- fee: fee total
- nota_precio: nota sobre precios
- pasos: lista de próximos pasos
- sobre_nosotros: breve descripción de la empresa (default sobre Nexus FinLabs)

Responde ÚNICAMENTE con el JSON, sin markdown ni explicaciones.""",
}


def load_env():
    """Load .env variables if running outside Docker."""
    for path in ['/home/albi_agent/openclawd_stack/.env', '.env']:
        if os.path.exists(path):
            with open(path) as f:
                for line in f:
                    if '=' in line and not line.startswith('#'):
                        k, v = line.strip().split('=', 1)
                        os.environ.setdefault(k, v.strip('"\''))


def expand_context_with_llm(doc_type, context):
    """Use LLM to expand a brief context into structured template variables."""
    api_key = os.environ.get("OPENROUTER_API_KEY", "")
    if not api_key:
        print("WARNING: No OPENROUTER_API_KEY, using basic context parsing")
        return {"raw_context": context}

    prompt_template = EXPANSION_PROMPTS.get(doc_type)
    if not prompt_template:
        return {"raw_context": context}

    prompt = prompt_template.format(context=context)

    try:
        resp = requests.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "openai/gpt-4o-mini",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.2,
                "max_tokens": 2000,
            },
            timeout=30,
        )

        if resp.status_code == 200:
            content = resp.json()["choices"][0]["message"]["content"].strip()
            # Strip markdown code fences if present
            if content.startswith("```"):
                content = content.split("\n", 1)[1]
                content = content.rsplit("```", 1)[0]
            data = json.loads(content)
            print(f"LLM expansion OK: {len(data)} fields")
            return data
        else:
            print(f"LLM API error {resp.status_code}: {resp.text[:200]}")

    except json.JSONDecodeError as e:
        print(f"LLM returned invalid JSON: {e}")
    except Exception as e:
        print(f"LLM expansion error: {e}")

    return {"raw_context": context}


def generate_pdf_from_template(doc_type, variables, output_path):
    """Generate PDF using Jinja2 template + WeasyPrint."""
    template_file = TEMPLATE_MAP.get(doc_type)
    env = Environment(loader=FileSystemLoader(TEMPLATES_DIR))

    try:
        template = env.get_template(template_file)
        html_content = template.render(**variables)
        HTML(string=html_content).write_pdf(output_path)
        return True
    except Exception as e:
        print(f"Template rendering error: {e}")
        return False


def generate_pdf_generic(doc_type, content, ref, output_path):
    """Fallback: generate generic PDF from plain text content."""
    html_content = f"""<!DOCTYPE html>
<html lang="es">
<head><meta charset="UTF-8"/>
<style>
    body {{ font-family: Arial, sans-serif; font-size: 14px; line-height: 1.4; color: #333; padding: 40px; }}
    h1 {{ text-align: center; font-size: 22px; text-transform: uppercase; font-weight: bold;
         border-bottom: 2px solid #333; padding-bottom: 10px; }}
    h2 {{ font-size: 16px; text-transform: uppercase; font-weight: bold;
         border-top: 1px solid #ccc; padding-top: 12px; margin-top: 24px; }}
    ul {{ margin: 6px 0 12px 24px; }}
    p {{ margin: 6px 0; }}
    .signature {{ margin-top: 40px; border-top: 1px solid #000; width: 300px; padding-top: 8px; }}
</style>
</head>
<body>
<h1>{doc_type} — Documento Generado</h1>
<p><strong>Fecha:</strong> {datetime.now().strftime('%d/%m/%Y')}</p>
<p><strong>Referencia:</strong> {ref}</p>
<hr/>
"""
    for line in content.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if (len(stripped) > 3 and stripped[0].isdigit() and '.' in stripped[:4]) or stripped.isupper():
            html_content += f"<h2>{stripped}</h2>\n"
        elif stripped.startswith('- ') or stripped.startswith('• '):
            html_content += f"<li>{stripped[2:]}</li>\n"
        else:
            html_content += f"<p>{stripped}</p>\n"

    html_content += """
<div class="signature">
<p>Firma: ______________________</p>
<p>Fecha: ______________________</p>
</div>
</body></html>"""

    HTML(string=html_content).write_pdf(output_path)


def create_docx_from_variables(variables, doc_type, output_path):
    """Create DOCX from template variables."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    title = doc.add_heading(f'{doc_type} — Documento Generado', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {variables.get("fecha", datetime.now().strftime("%d/%m/%Y"))}')
    doc.add_paragraph(f'Referencia: {variables.get("ref", "")}')
    doc.add_paragraph('')

    # Iterate over variables and add them as structured sections
    skip_keys = {'ref', 'fecha', 'raw_context', 'title'}

    for key, value in variables.items():
        if key in skip_keys:
            continue

        # Section header
        header = key.replace('_', ' ').title()

        if isinstance(value, str) and value:
            doc.add_heading(header, level=2)
            doc.add_paragraph(value)
        elif isinstance(value, list):
            doc.add_heading(header, level=2)
            for item in value:
                if isinstance(item, dict):
                    parts = [f"{v}" for k, v in item.items() if v]
                    doc.add_paragraph(" — ".join(parts), style='List Bullet')
                elif isinstance(item, str):
                    doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('Firma:')
    doc.save(output_path)


def create_docx_from_text(text, doc_type, output_path):
    """Fallback: create DOCX from plain text."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    title = doc.add_heading(f'{doc_type} — Documento Generado', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph('')

    for line in text.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if (len(stripped) > 3 and stripped[0].isdigit() and '.' in stripped[:4]) or stripped.isupper():
            doc.add_heading(stripped, level=2)
        elif stripped.startswith('- ') or stripped.startswith('• '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            doc.add_paragraph(stripped)

    doc.add_paragraph('')
    doc.add_paragraph('_' * 40)
    doc.add_paragraph('Firma:')
    doc.save(output_path)


def main():
    if len(sys.argv) < 3:
        print("Usage: python document_generator_cli.py <DOC_TYPE> <CONTEXT>")
        print("Types with templates: NDA, SOW, PROPUESTA")
        print("Other types: generic document generation")
        sys.exit(1)

    load_env()

    doc_type = sys.argv[1].upper()
    context = sys.argv[2]

    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    ref = str(uuid.uuid4())[:6].upper()
    base_name = f"{doc_type}_{timestamp}_{ref}"

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    print(f"--- INICIANDO GENERACION DE DOCUMENTO ---")
    print(f"Tipo: {doc_type}")
    print(f"Ref: {base_name}")

    pdf_path = os.path.join(OUTPUT_DIR, f"{base_name}.pdf")
    docx_path = os.path.join(OUTPUT_DIR, f"{base_name}.docx")

    has_template = doc_type in TEMPLATE_MAP

    if has_template:
        # ── Template-based generation ──
        print(f"Template encontrada: {TEMPLATE_MAP[doc_type]}")
        print("Expandiendo contexto con LLM...")

        variables = expand_context_with_llm(doc_type, context)
        variables["ref"] = base_name
        variables["fecha"] = datetime.now().strftime("%d/%m/%Y")

        # Alias for PROPOSAL -> PROPUESTA
        if doc_type == "PROPOSAL":
            doc_type = "PROPUESTA"

        print("Generando PDF desde template...")
        success = generate_pdf_from_template(doc_type, variables, pdf_path)
        if not success:
            print("Template fallback: generando PDF genérico...")
            generate_pdf_generic(doc_type, context, base_name, pdf_path)

        print(f"PDF_FILE:{base_name}.pdf")

        print("Generando DOCX...")
        create_docx_from_variables(variables, doc_type, docx_path)
        print(f"DOCX_FILE:{base_name}.docx")

    else:
        # ── Generic generation (original behavior) ──
        print("Sin template específica. Generando documento genérico...")

        print("Generando PDF...")
        generate_pdf_generic(doc_type, context, base_name, pdf_path)
        print(f"PDF_FILE:{base_name}.pdf")

        print("Generando DOCX...")
        create_docx_from_text(context, doc_type, docx_path)
        print(f"DOCX_FILE:{base_name}.docx")

    print("--------------------------------------------------")
    print(f"SUCCESS: Documentos generados correctamente.")
    print(f"PDF: {pdf_path}")
    print(f"DOCX: {docx_path}")


if __name__ == "__main__":
    main()
